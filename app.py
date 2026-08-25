"""Chat RAG sobre un convenio colectivo en PDF con Streamlit y LangChain."""

import os
import tempfile
import io
import datetime

import streamlit as st
import pandas as pd
from docx import Document
from langchain_classic.chains import RetrievalQA
from langchain_core.prompts import PromptTemplate
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Mensajes y URLs
NO_INFO_MESSAGE = "No encuentro esta información en el convenio subido"
REGCON_SEARCH_URL = "https://expinterweb.mites.gob.es/regcon/pub/buscadorTextosEstatal?language=es"

# Prompt actualizado para RAG
RAG_PROMPT = PromptTemplate(
    input_variables=["context", "question"],
    template=f"""Eres un graduado social experto en convenios colectivos.

Lee el siguiente contexto del convenio y responde a la pregunta del usuario.
Normas:
1. Basa tu respuesta ÚNICAMENTE en el contexto proporcionado.
2. Puedes resumir o explicar con palabras sencillas, pero no inventes datos.
3. Si la pregunta es muy general, haz un resumen de todo lo que encuentres.
4. Si hay tablas de salarios o categorías, numéralas literalmente.
5. Si no se habla de lo que pide el usuario, responde exactamente: {NO_INFO_MESSAGE}

Contexto del convenio:
{{context}}

Pregunta: {{question}}

Respuesta:""",
)

def reset_document_state() -> None:
    st.session_state.vectorstore = None
    st.session_state.pdf_name = None
    st.session_state.messages = []

def initialise_session_state() -> None:
    if "vectorstore" not in st.session_state:
        st.session_state.vectorstore = None
    if "pdf_name" not in st.session_state:
        st.session_state.pdf_name = None
    if "messages" not in st.session_state:
        st.session_state.messages = []

@st.cache_resource(show_spinner=False)
def get_local_embeddings() -> HuggingFaceEmbeddings:
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

def build_vectorstore(uploaded_pdf) -> FAISS:
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(uploaded_pdf.getvalue())
            temp_path = temp_file.name

        documents = PyPDFLoader(temp_path).load()
        if not documents:
            raise ValueError("El PDF no contiene texto procesable.")

        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_documents(documents)
        if not chunks:
            raise ValueError("No se han podido crear fragmentos.")

        embeddings = get_local_embeddings()
        return FAISS.from_documents(chunks, embeddings)
    finally:
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)

def create_qa_chain(vectorstore: FAISS, gemini_api_key: str) -> RetrievalQA:
    llm = ChatGoogleGenerativeAI(
        model="gemini-3.6-flash", 
        temperature=0,
        api_key=gemini_api_key,
    )
    return RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 15}),
        chain_type_kwargs={"prompt": RAG_PROMPT},
        return_source_documents=True,
    )

def generar_word_baja(empresa, trabajador, dni, fecha, motivo):
    """Genera un archivo Word en memoria según el tipo de baja"""
    doc = Document()
    doc.add_heading('COMUNICACIÓN DE CESE', 0)
    doc.add_paragraph(f"Empresa: {empresa}")
    doc.add_paragraph(f"A la atención de D./Dña: {trabajador}")
    doc.add_paragraph(f"DNI/NIE: {dni}\n")
    
    # Textos de ejemplo (puedes cambiarlos por vuestras plantillas reales)
    if motivo == "Baja Voluntaria":
        texto = f"Por la presente, la dirección de la empresa acusa recibo de su comunicación en la que nos informa de su decisión de causar baja voluntaria en la empresa con efectos del día {fecha}.\n\nAceptamos su decisión y le informamos que tiene a su disposición en nuestras oficinas su liquidación y finiquito."
    elif motivo == "Fin de Contrato":
        texto = f"Por la presente le comunicamos que, con fecha {fecha}, finalizará su contrato de trabajo por expiración del tiempo convenido, procediendo a su baja en la Seguridad Social.\n\nLe agradecemos los servicios prestados y le informamos que su liquidación estará a su disposición."
    elif motivo == "No superación periodo de prueba":
        texto = f"Por la presente le comunicamos la decisión de la empresa de dar por finalizada su relación laboral por no superación del periodo de prueba, con efectos del día {fecha}."
    else:
        texto = f"Se le comunica que con fecha {fecha} se procede a su baja en la empresa por el siguiente motivo: {motivo}."
        
    doc.add_paragraph(texto)
    doc.add_paragraph("\n\n\nFdo. La Empresa                                   Fdo. El Trabajador (Recibí)")
    
    # Guardar en memoria
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def apply_custom_style() -> None:
    st.markdown(
        """
        <style>
            :root { --ink: #172033; --muted: #64748b; --surface: #ffffff; --line: #dbe4ef; --blue: #2563eb; }
            .stApp { background: radial-gradient(circle at 88% 4%, #dceeff 0, transparent 26rem), linear-gradient(180deg, #f8fbff 0%, #f2f6fb 100%); color: var(--ink); }
            .block-container { max-width: 1080px; padding-top: 2rem; padding-bottom: 3rem; }
            section[data-testid="stSidebar"] { background: linear-gradient(180deg, #102f57 0%, #153b6d 100%); }
            section[data-testid="stSidebar"] * { color: #f8fbff; }
            section[data-testid="stSidebar"] input, section[data-testid="stSidebar"] select { background: rgba(255, 255, 255, 0.96) !important; color: var(--ink) !important; border-radius: 10px !important; }
            .app-hero { background: linear-gradient(135deg, #ffffff 0%, #edf6ff 100%); border: 1px solid rgba(37, 99, 235, 0.14); border-radius: 22px; padding: 2rem; margin-bottom: 1rem; }
            .app-hero h1 { margin-top: 0; }
        </style>
        """,
        unsafe_allow_html=True,
    )

def main() -> None:
    st.set_page_config(page_title="Portal Laboral", page_icon="⚖️", layout="wide")
    initialise_session_state()
    apply_custom_style()

    st.markdown(
        """
        <section class="app-hero">
            <div style="color: #2563eb; font-weight: 800; font-size: 0.8rem; text-transform: uppercase;">Portal Laboral · IA & Plantillas</div>
            <h1>Herramientas del Departamento</h1>
        </section>
        """,
        unsafe_allow_html=True,
    )

    # --- BARRA LATERAL ---
    with st.sidebar:
        st.markdown("<h2>⚙️ Configuración</h2>", unsafe_allow_html=True)
        gemini_api_key = st.text_input("API Key de Gemini", type="password")
        
        st.divider()
        st.subheader("🔍 Buscador de Convenios")
        
        datos_convenios = {
            "Sector": ["Hostelería", "Metal", "Oficinas y Despachos"],
            "Provincia": ["Ourense", "Estatal", "Madrid"],
            "Enlace": [
                "https://www.boe.es/boe/dias/2021/11/17/pdfs/BOE-A-2021-18894.pdf",
                "https://www.boe.es/boe/dias/2022/01/12/pdfs/BOE-A-2022-478.pdf", 
                "https://www.boe.es/boe/dias/2023/12/28/pdfs/BOE-A-2023-26462.pdf" 
            ]
        }
        df_convenios = pd.DataFrame(datos_convenios)
        sector_elegido = st.selectbox("1. Sector:", df_convenios["Sector"].unique())
        provincias_disponibles = df_convenios[df_convenios["Sector"] == sector_elegido]["Provincia"].unique()
        provincia_elegida = st.selectbox("2. Provincia:", provincias_disponibles)
        
        resultado = df_convenios[(df_convenios["Sector"] == sector_elegido) & (df_convenios["Provincia"] == provincia_elegida)]
        if not resultado.empty:
            st.link_button("⬇️ Descargar PDF Oficial", resultado.iloc[0]["Enlace"], use_container_width=True)

        st.link_button("🌐 Buscar en REGCON ↗", REGCON_SEARCH_URL, use_container_width=True)

    # --- PESTAÑAS PRINCIPALES ---
    tab1, tab2 = st.tabs(["📚 Asistente de Convenios (Chat RAG)", "📝 Generador de Plantillas (Bajas)"])

    # --- PESTAÑA 1: CHAT CONVENIOS ---
    with tab1:
        st.subheader("Sube un convenio para analizarlo")
        uploaded_pdf = st.file_uploader("Sube el convenio en PDF", type=["pdf"])
        
        if uploaded_pdf and uploaded_pdf.name != st.session_state.pdf_name:
            reset_document_state()

        if uploaded_pdf and st.session_state.vectorstore is None:
            try:
                with st.spinner("Leyendo y preparando el convenio..."):
                    st.session_state.vectorstore = build_vectorstore(uploaded_pdf)
                    st.session_state.pdf_name = uploaded_pdf.name
                st.success(f"¡Listo! Convenio {uploaded_pdf.name} procesado.")
            except Exception as error:
                st.error(f"Error al procesar: {error}")

        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        question = st.chat_input("Pregunta sobre el convenio...", disabled=st.session_state.vectorstore is None or not gemini_api_key)

        if st.session_state.vectorstore is not None and not gemini_api_key:
            st.warning("⚠️ Introduce tu API Key de Gemini en la barra lateral izquierda para chatear.")

        if question:
            st.session_state.messages.append({"role": "user", "content": question})
            with st.chat_message("user"):
                st.markdown(question)

            with st.chat_message("assistant"):
                try:
                    with st.spinner("Buscando respuesta..."):
                        qa_chain = create_qa_chain(st.session_state.vectorstore, gemini_api_key)
                        response = qa_chain.invoke({"query": question})
                        answer = response["result"]
                    st.markdown(answer)
                except Exception as error:
                    answer = f"Error: {error}"
                    st.error(answer)

            st.session_state.messages.append({"role": "assistant", "content": answer})

    # --- PESTAÑA 2: GENERADOR DE BAJAS ---
    with tab2:
        st.subheader("Generador rápido de comunicaciones")
        st.write("Rellena los datos y descarga la carta de baja o fin de contrato en formato Word (lista para firmar).")
        
        with st.form("form_bajas"):
            colA, colB = st.columns(2)
            with colA:
                input_empresa = st.text_input("Razón Social de la Empresa")
                input_trabajador = st.text_input("Nombre y Apellidos del Trabajador")
                input_dni = st.text_input("DNI / NIE")
            with colB:
                input_fecha = st.date_input("Fecha de Efecto", format="DD/MM/YYYY")
                # Formatear la fecha para el documento
                fecha_str = input_fecha.strftime("%d/%m/%Y")
                input_motivo = st.selectbox(
                    "Motivo de la Baja", 
                    ["Baja Voluntaria", "Fin de Contrato", "No superación periodo de prueba"]
                )
            
            btn_generar = st.form_submit_button("Redactar Documento", use_container_width=True)
            
        if btn_generar:
            if not input_empresa or not input_trabajador or not input_dni:
                st.error("Por favor, rellena todos los datos del formulario.")
            else:
                # Llama a la función que crea el Word
                word_file = generar_word_baja(input_empresa, input_trabajador, input_dni, fecha_str, input_motivo)
                st.success("¡Documento generado con éxito!")
                
                # Botón nativo de Streamlit para descargar
                st.download_button(
                    label="⬇️ Descargar Word (.docx)",
                    data=word_file,
                    file_name=f"Comunicacion_Baja_{input_trabajador.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

if __name__ == "__main__":
    main()
